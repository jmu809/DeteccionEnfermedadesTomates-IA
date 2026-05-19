"""Convierte PROYECTO.md a PROYECTO.docx con formato académico."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_inline_formatting(paragraph, text):
    """Procesa **negrita**, *cursiva* y `código` dentro de un párrafo."""
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            paragraph.add_run(text[last:m.start()])
        if m.group(2):          # **negrita**
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(3):        # *cursiva*
            run = paragraph.add_run(m.group(3))
            run.italic = True
        elif m.group(4):        # `código`
            run = paragraph.add_run(m.group(4))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        last = m.end()
    if last < len(text):
        paragraph.add_run(text[last:])


def clean(text):
    """Elimina anclas HTML <a name=...> y etiquetas img del texto."""
    text = re.sub(r'<a name="[^"]*"></a>', '', text)
    text = re.sub(r'<[^>]+>', '', text)
    return text.strip()


def add_heading(doc, text, level):
    text = clean(text)
    if not text:
        return
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.runs[0].font.size = Pt(18)
        p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        p.runs[0].font.size = Pt(14)
        p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    elif level == 3:
        p.runs[0].font.size = Pt(12)
        p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def add_paragraph(doc, text, style='Normal'):
    text = clean(text)
    if not text:
        return
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    add_inline_formatting(p, text)


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)
    doc.add_paragraph()


def add_table(doc, header, rows):
    n_cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Cabecera
    for i, cell_text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = cell_text.strip()
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, '2E74B5')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Filas
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            p = cell.paragraphs[0]
            add_inline_formatting(p, cell_text.strip())
            p.runs[0].font.size = Pt(9) if p.runs else None
            if r_idx % 2 == 0:
                set_cell_bg(cell, 'DEEAF1')

    doc.add_paragraph()


# ── Parser principal ───────────────────────────────────────────────────────────

def parse_md_to_docx(md_path: str, docx_path: str):
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # Estilo base
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    lines = Path(md_path).read_text(encoding='utf-8').splitlines()
    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # Bloque de código
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                add_code_block(doc, code_lines)
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Tabla Markdown
        if '|' in line and i + 1 < len(lines) and re.match(r'^\|[-| :]+\|', lines[i + 1]):
            header = [c for c in line.split('|') if c.strip()]
            i += 2  # saltar separador
            rows = []
            while i < len(lines) and '|' in lines[i]:
                row = [c for c in lines[i].split('|') if c.strip() != '']
                if row:
                    rows.append(row)
                i += 1
            add_table(doc, header, rows)
            continue

        # Headings
        if line.startswith('#### '):
            add_heading(doc, line[5:], 4)
        elif line.startswith('### '):
            add_heading(doc, line[4:], 3)
        elif line.startswith('## '):
            add_heading(doc, line[3:], 2)
        elif line.startswith('# '):
            add_heading(doc, line[2:], 1)

        # Separador horizontal
        elif re.match(r'^---+$', line.strip()):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '2E74B5')
            pBdr.append(bottom)
            pPr.append(pBdr)

        # Imagen (mostrar como texto)
        elif line.strip().startswith('!['):
            m = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
            if m:
                p = doc.add_paragraph()
                run = p.add_run(f'[Figura: {m.group(1)}  —  {m.group(2)}]')
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Lista con viñeta (- o *)
        elif re.match(r'^(\s*)([-*]) ', line):
            indent = len(line) - len(line.lstrip())
            text = re.sub(r'^(\s*)[-*] ', '', line)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(1 + indent * 0.3)
            p.paragraph_format.space_after = Pt(2)
            add_inline_formatting(p, clean(text))

        # Lista numerada
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(2)
            add_inline_formatting(p, clean(text))

        # Línea en blanco
        elif line.strip() == '':
            pass

        # Párrafo normal
        else:
            text = clean(line)
            if text:
                p = doc.add_paragraph(style='Normal')
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.first_line_indent = Cm(0)
                add_inline_formatting(p, text)

        i += 1

    doc.save(docx_path)
    print(f'Documento guardado: {docx_path}')


if __name__ == '__main__':
    parse_md_to_docx(
        r'c:\Users\jrmu7\OneDrive\Desktop\Uni\IA-tomates\PROYECTO.md',
        r'c:\Users\jrmu7\OneDrive\Desktop\Uni\IA-tomates\PROYECTO.docx'
    )
